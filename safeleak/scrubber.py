import io
import logging
from typing import Tuple

import fitz
from docx import Document
from PIL import Image
from presidio_analyzer import AnalyzerEngine

logger = logging.getLogger(__name__)

_analyzer: AnalyzerEngine = None


def _get_analyzer() -> AnalyzerEngine:
    global _analyzer
    if _analyzer is None:
        logger.info("Loading Presidio AnalyzerEngine + spaCy en_core_web_lg...")
        _analyzer = AnalyzerEngine()
        logger.info("Presidio + spaCy loaded")
    return _analyzer


class ScrubberAgent:
    def __init__(self):
        self.analyzer = _get_analyzer()
        logger.info("ScrubberAgent ready")

    def scrub(self, file_bytes: bytes, filename: str) -> dict:
        ext = filename.rsplit(".", 1)[-1].lower()
        metadata_removed = []
        clean_bytes = file_bytes

        try:
            if ext == "pdf":
                try:
                    clean_bytes, metadata_removed = self._strip_pdf_metadata(file_bytes)
                except Exception as e:
                    return {"error": f"Could not parse PDF: {e}. File may be corrupted, password-protected, or not a valid PDF."}
            elif ext == "docx":
                try:
                    clean_bytes, metadata_removed = self._strip_docx_metadata(file_bytes)
                except Exception as e:
                    return {"error": f"Could not parse DOCX: {e}. File may be corrupted or not a valid Word document."}
            elif ext in ("jpg", "jpeg", "png"):
                try:
                    clean_bytes, metadata_removed = self._strip_image_metadata(file_bytes)
                except Exception as e:
                    return {"error": f"Could not parse image: {e}. File may be corrupted or not a valid image."}
            elif ext == "txt":
                clean_bytes = file_bytes
                metadata_removed = []
            else:
                return {"error": f"Unsupported file type: {ext}"}

            text = self._extract_text(clean_bytes, filename)
            presidio_pii = self._run_pii_detection(text)

            # Gemini second-pass PII analysis — always returns a dict, never raises
            from gemini_pii import analyse_pii_with_gemini
            gemini_result = analyse_pii_with_gemini(text)
            gemini_findings = gemini_result.get("findings", [])
            gemini_status = gemini_result.get("status", "unavailable")

            # Merge results (Presidio first, then Gemini additions)
            pii_found = presidio_pii + gemini_findings

            return {
                "clean_bytes": clean_bytes,
                "metadata_removed": metadata_removed,
                "pii_found": pii_found,
                "pii_count": len(pii_found),
                "gemini_status": gemini_status,
            }
        except Exception as e:
            logger.exception("Scrub error")
            return {"error": str(e)}

    def _strip_pdf_metadata(self, file_bytes: bytes) -> Tuple[bytes, list]:
        removed = []
        buf = io.BytesIO(file_bytes)
        doc = fitz.open(stream=buf, filetype="pdf")

        meta = doc.metadata or {}
        field_names = ["author", "creator", "producer", "subject", "title", "keywords", "trapped"]
        for field in field_names:
            val = meta.get(field, "")
            if val:
                removed.append(f"{field.capitalize()}: {val}")

        doc.set_metadata({})

        out = io.BytesIO()
        doc.save(out, garbage=4, deflate=True, clean=True)
        doc.close()
        return out.getvalue(), removed

    def _strip_docx_metadata(self, file_bytes: bytes) -> Tuple[bytes, list]:
        removed = []
        buf = io.BytesIO(file_bytes)
        doc = Document(buf)
        props = doc.core_properties

        # Core properties (writable via python-docx)
        core_fields = {
            "author": "Author",
            "last_modified_by": "Last Modified By",
            "title": "Title",
            "subject": "Subject",
            "keywords": "Keywords",
            "description": "Description",
            "category": "Category",
        }

        for attr, label in core_fields.items():
            try:
                val = getattr(props, attr, None)
                if val:
                    removed.append(f"{label}: {val}")
                    setattr(props, attr, "")
            except Exception:
                pass

        # Save the docx to a buffer first, then patch app.xml directly via zipfile
        out = io.BytesIO()
        doc.save(out)
        docx_bytes = out.getvalue()

        # Strip app properties (company, manager, etc.) from docProps/app.xml in the zip
        try:
            import zipfile
            from lxml import etree as _etree

            APP_XML = "docProps/app.xml"
            APP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"
            STRIP_TAGS = ("Company", "Manager", "Application", "Template", "HyperlinkBase")

            new_zip_buf = io.BytesIO()
            with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zin, \
                 zipfile.ZipFile(new_zip_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == APP_XML:
                        tree = _etree.fromstring(data)
                        for tag in STRIP_TAGS:
                            el = tree.find(f"{{{APP_NS}}}{tag}")
                            if el is not None and el.text and el.text.strip():
                                removed.append(f"{tag}: {el.text.strip()}")
                                el.text = ""
                        data = _etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
                    zout.writestr(item, data)

            return new_zip_buf.getvalue(), removed
        except Exception as e:
            logger.debug(f"App properties zip strip skipped: {e}")
            return docx_bytes, removed

    def _strip_image_metadata(self, file_bytes: bytes) -> Tuple[bytes, list]:
        removed = []
        buf = io.BytesIO(file_bytes)
        img = Image.open(buf)

        try:
            exif_data = img._getexif() or {}
            from PIL.ExifTags import TAGS
            tag_map = {v: k for k, v in TAGS.items()}

            interest = {
                "GPSInfo": "GPS",
                "Make": "Camera Make",
                "Model": "Camera Model",
                "Software": "Software",
                "Artist": "Artist",
                "Copyright": "Copyright",
                "DateTime": "DateTime",
                "DateTimeOriginal": "DateTimeOriginal",
                "LensModel": "Lens Model",
            }

            for tag_name, label in interest.items():
                tag_id = tag_map.get(tag_name)
                if tag_id and tag_id in exif_data:
                    val = exif_data[tag_id]
                    if tag_name == "GPSInfo" and isinstance(val, dict):
                        removed.append("GPS: (coordinates stripped)")
                    else:
                        removed.append(f"{label}: {val}")
        except Exception:
            pass

        img_no_exif = Image.new(img.mode, img.size)
        img_no_exif.putdata(list(img.getdata()))

        out = io.BytesIO()
        fmt = img.format or "PNG"
        if fmt == "JPEG":
            img_no_exif.save(out, format="JPEG", quality=95)
        else:
            img_no_exif.save(out, format="PNG")
        return out.getvalue(), removed

    def _extract_text(self, clean_bytes: bytes, filename: str) -> str:
        ext = filename.rsplit(".", 1)[-1].lower()
        try:
            if ext == "pdf":
                buf = io.BytesIO(clean_bytes)
                doc = fitz.open(stream=buf, filetype="pdf")
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            elif ext == "docx":
                buf = io.BytesIO(clean_bytes)
                doc = Document(buf)
                return "\n".join(p.text for p in doc.paragraphs)
            elif ext == "txt":
                return clean_bytes.decode("utf-8", errors="replace")
            else:
                return ""
        except Exception:
            return ""

    def _run_pii_detection(self, text: str) -> list:
        if not text.strip():
            return []

        entities = [
            "PERSON",
            "EMAIL_ADDRESS",
            "PHONE_NUMBER",
            "LOCATION",
            "DATE_TIME",
            "IBAN_CODE",
            "CREDIT_CARD",
            "CRYPTO",
            "IP_ADDRESS",
            "NRP",
            "MEDICAL_LICENSE",
        ]

        try:
            results = self.analyzer.analyze(
                text=text,
                entities=entities,
                language="en",
                score_threshold=0.6,
            )
        except Exception as e:
            logger.error(f"PII detection error: {e}")
            return []

        pii_list = []
        for r in results:
            pii_list.append({
                "type": r.entity_type,
                "text": "[REDACTED]",
                "score": round(r.score, 3),
                "method": "presidio",
            })

        return pii_list
